from datetime import timedelta
from decimal import Decimal
from io import BytesIO

from django.contrib.auth.models import User
from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document

from .documentos_banca import (
    gerar_docx_ata,
    montar_dados_ata,
)
from .models import (
    BancaTCC,
    ComposicaoBanca,
    Discente,
    EspacoFisico,
    ProjetoTCC,
    SolicitacaoAgendamento,
    pUsuario,
)


class DocumentosInstitucionaisTests(TestCase):

    def setUp(self):

        self.usuario_coordenacao = User.objects.create_user(
            username='coordenacao.ata@ufac.br',
            password='Senha123!',
            is_active=True
        )

        pUsuario.objects.create(
            usuario=self.usuario_coordenacao,
            perfil='COORDENACAO'
        )

        self.usuario_orientador = User.objects.create_user(
            username='orientador.ata@ufac.br',
            password='Senha123!',
            first_name='Maria',
            last_name='Silva',
            is_active=True
        )

        self.orientador = pUsuario.objects.create(
            usuario=self.usuario_orientador,
            perfil='DOCENTE',
            titulacao='PROFA_DRA'
        )

        self.usuario_avaliador = User.objects.create_user(
            username='avaliador.ata@ufac.br',
            password='Senha123!',
            first_name='João',
            last_name='Souza',
            is_active=True
        )

        self.avaliador = pUsuario.objects.create(
            usuario=self.usuario_avaliador,
            perfil='DOCENTE',
            titulacao='PROF_ME'
        )

        self.espaco = EspacoFisico.objects.create(
            nome='Laboratório Web Academy'
        )

        self.discente = Discente.objects.create(
            nome='Ana Discente',
            matricula='20260000999'
        )

        self.projeto = ProjetoTCC.objects.create(
            titulo='Sistema de Gestão de Bancas de TCC',
            resumo='Resumo do trabalho.',
            semestre_letivo='2026.2',
            discente=self.discente,
            status='APROVADO'
        )

        inicio = (
            timezone.now()
            + timedelta(days=10)
        ).replace(
            hour=14,
            minute=30,
            second=0,
            microsecond=0
        )

        fim = inicio + timedelta(hours=2)

        self.solicitacao = SolicitacaoAgendamento.objects.create(
            usuario_solicitante=self.orientador,
            projeto_tcc=self.projeto,
            espaco=self.espaco,
            opcao_data_inicio=inicio,
            opcao_data_fim=fim,
            status='APROVADA'
        )

        self.composicao = ComposicaoBanca.objects.create(
            projeto_tcc=self.projeto,
            solicitacao=self.solicitacao,
            orientador=self.orientador,
            avaliador_interno=self.avaliador,
            presidente=self.orientador,
            nome_avaliador_externo='Carlos Externo',
            titulacao_avaliador_externo='PROF_DR',
            instituicao_avaliador_externo='IFAC'
        )

        self.banca = BancaTCC.objects.create(
            solicitacao=self.solicitacao,
            projeto_tcc=self.projeto,
            espaco=self.espaco,
            data_horario_inicio=inicio,
            data_horario_fim=fim,
            status='AGENDADA'
        )

    def _dados(self):

        return montar_dados_ata(
            self.solicitacao,
            self.composicao,
            self.banca
        )

    def _texto_docx(self, arquivo):

        documento = Document(arquivo)

        textos = [
            paragrafo.text
            for paragrafo in documento.paragraphs
        ]

        for tabela in documento.tables:
            for linha in tabela.rows:
                textos.extend(
                    celula.text
                    for celula in linha.cells
                )

        return '\n'.join(textos)

    def test_presidente_aparece_na_funcao_do_integrante(self):

        dados = self._dados()

        integrantes_orientadores = [
            integrante
            for integrante in dados['integrantes']
            if 'Maria Silva' in integrante['nome']
        ]

        self.assertEqual(
            len(integrantes_orientadores),
            1
        )

        self.assertEqual(
            integrantes_orientadores[0]['funcao'],
            'Orientador(a) e presidente'
        )

        integrante_externo = next(
            integrante
            for integrante in dados['integrantes']
            if 'Carlos Externo' in integrante['nome']
        )

        self.assertEqual(
            integrante_externo['funcao'],
            'Avaliador(a) externo(a)'
        )

        self.assertEqual(
            integrante_externo['instituicao'],
            'IFAC'
        )

    def test_docx_pre_defesa_mantem_campos_em_branco(self):

        arquivo = gerar_docx_ata(self._dados())
        documento = Document(arquivo)

        texto = '\n'.join(
            paragrafo.text
            for paragrafo in documento.paragraphs
        )

        texto_negrito = ''.join(
            run.text
            for paragrafo in documento.paragraphs
            for run in paragrafo.runs
            if run.bold
        )

        self.assertIn(
            'ATA DE APRESENTAÇÃO DO TRABALHO DE CONCLUSÃO DE CURSO',
            texto
        )

        self.assertIn(
            'Resultado: ______________________________',
            texto
        )

        self.assertIn(
            'Nota: __________________',
            texto
        )

        self.assertIn(
            'Ana Discente',
            texto_negrito
        )

        self.assertIn(
            'Sistema de Gestão de Bancas de TCC',
            texto_negrito
        )

        self.assertIn(
            self._dados()['data_defesa_extenso'],
            texto_negrito
        )

        self.assertIn(
            'Instituição: IFAC',
            texto
        )

        self.assertNotIn('Código CRC', texto)
        self.assertNotIn('Documento assinado eletronicamente', texto)

    def test_docx_final_exibe_nota_e_prazo(self):

        self.banca.status = 'FINALIZADA'
        self.banca.nota = Decimal('9.75')

        self.banca.save(
            update_fields=['status', 'nota']
        )

        texto = self._texto_docx(
            gerar_docx_ata(self._dados())
        )

        self.assertIn(
            '9,75 (nove vírgula setenta e cinco)',
            texto
        )

        self.assertIn(
            '30 (trinta) dias corridos',
            texto
        )

        self.assertNotIn(
            'Nota: __________________',
            texto
        )

    def test_coordenacao_gera_pdf_da_ata(self):

        self.client.force_login(
            self.usuario_coordenacao
        )

        response = self.client.get(
            reverse(
                'gerar_pdf_banca',
                args=[self.solicitacao.id]
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/pdf'
        )
        self.assertTrue(response.content.startswith(b'%PDF'))
        self.assertIn(
            'pre_defesa.pdf',
            response['Content-Disposition']
        )

    def test_coordenacao_gera_docx_da_ata(self):

        self.client.force_login(
            self.usuario_coordenacao
        )

        response = self.client.get(
            reverse(
                'gerar_docx_banca',
                args=[self.solicitacao.id]
            )
        )

        conteudo = b''.join(
            response.streaming_content
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(conteudo.startswith(b'PK'))
        self.assertIn(
            'pre_defesa.docx',
            response['Content-Disposition']
        )

        texto = self._texto_docx(
            BytesIO(conteudo)
        )

        self.assertIn('Ana Discente', texto)
        self.assertIn('Prof. Dr. Carlos Externo', texto)
        self.assertIn('IFAC', texto)

    def test_docente_sem_vinculo_nao_gera_docx(self):

        usuario_sem_vinculo = User.objects.create_user(
            username='sem.vinculo.ata@ufac.br',
            password='Senha123!',
            is_active=True
        )

        pUsuario.objects.create(
            usuario=usuario_sem_vinculo,
            perfil='DOCENTE'
        )

        self.client.force_login(usuario_sem_vinculo)

        response = self.client.get(
            reverse(
                'gerar_docx_banca',
                args=[self.solicitacao.id]
            )
        )

        self.assertRedirects(
            response,
            reverse('documentos')
        )

    def test_documentos_oferece_pdf_e_docx(self):

        self.client.force_login(
            self.usuario_orientador
        )

        response = self.client.get(
            reverse('documentos')
        )

        self.assertContains(
            response,
            reverse(
                'gerar_pdf_banca',
                args=[self.solicitacao.id]
            )
        )

        self.assertContains(
            response,
            reverse(
                'gerar_docx_banca',
                args=[self.solicitacao.id]
            )
        )

        self.assertContains(response, 'Pré-defesa')


class NotificacaoGlobalAvaliacaoTests(TestCase):

    def setUp(self):

        self.usuario_coordenacao = User.objects.create_user(
            username='coordenacao.alerta@ufac.br',
            password='Senha123!',
            is_active=True
        )

        pUsuario.objects.create(
            usuario=self.usuario_coordenacao,
            perfil='COORDENACAO'
        )

        usuario_orientador = User.objects.create_user(
            username='orientador.alerta@ufac.br',
            password='Senha123!',
            is_active=True
        )

        orientador = pUsuario.objects.create(
            usuario=usuario_orientador,
            perfil='DOCENTE'
        )

        usuario_avaliador = User.objects.create_user(
            username='avaliador.alerta@ufac.br',
            password='Senha123!',
            is_active=True
        )

        avaliador = pUsuario.objects.create(
            usuario=usuario_avaliador,
            perfil='DOCENTE'
        )

        espaco = EspacoFisico.objects.create(
            nome='Sala de avaliação do alerta'
        )

        discente = Discente.objects.create(
            nome='Discente do alerta',
            matricula='20260000888'
        )

        projeto = ProjetoTCC.objects.create(
            titulo='TCC do alerta global',
            resumo='Resumo.',
            semestre_letivo='2026.2',
            discente=discente
        )

        inicio = timezone.now() + timedelta(days=15)

        self.solicitacao = SolicitacaoAgendamento.objects.create(
            usuario_solicitante=orientador,
            projeto_tcc=projeto,
            espaco=espaco,
            opcao_data_inicio=inicio,
            opcao_data_fim=(
                inicio + timedelta(hours=2)
            ),
            status='EM_ANÁLISE'
        )

        ComposicaoBanca.objects.create(
            projeto_tcc=projeto,
            solicitacao=self.solicitacao,
            orientador=orientador,
            avaliador_interno=avaliador,
            presidente=None
        )

    def test_presidente_ausente_usa_notificacao_global(self):

        self.client.force_login(
            self.usuario_coordenacao
        )

        response = self.client.post(
            reverse(
                'avaliar_solicitacao',
                args=[self.solicitacao.id]
            ),
            {
                'acao': 'aprovar',
                'presidente': '',
                'motivo_decisao': 'Aprovação de teste.',
            }
        )

        self.assertEqual(response.status_code, 200)

        self.assertContains(
            response,
            'class="alerta alerta-error"'
        )

        self.assertContains(
            response,
            'Selecione o presidente da banca antes '
            'de aprovar a solicitação.'
        )

        self.assertNotContains(
            response,
            'class="errorlist"'
        )
