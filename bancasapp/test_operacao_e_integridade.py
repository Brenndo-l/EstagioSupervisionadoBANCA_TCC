import json
import sqlite3
from datetime import timedelta
from io import StringIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch
from zipfile import ZipFile

from django.contrib.auth.models import User
from django.core.management import call_command
from django.core.management.base import CommandError
from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document

from .documentos_banca import (
    gerar_docx_ata,
    montar_dados_ata,
)
from .models import (
    BancaTCC,
    ComposicaoBanca,
    Discente,
    EspacoFisico,
    ProjetoTCC,
    SolicitacaoAgendamento,
    pUsuario,
)


class BaseIntegridadeTestCase(TestCase):

    def setUp(self):

        self.usuario_coordenacao = User.objects.create_user(
            username='coordenacao.integridade@ufac.br',
            password='Senha123!',
            is_active=True,
        )

        pUsuario.objects.create(
            usuario=self.usuario_coordenacao,
            perfil='COORDENACAO',
        )

        usuario_orientador = User.objects.create_user(
            username='orientador.integridade@ufac.br',
            password='Senha123!',
            first_name='Maria',
            last_name='Orientadora',
            is_active=True,
        )

        self.orientador = pUsuario.objects.create(
            usuario=usuario_orientador,
            perfil='DOCENTE',
            titulacao='PROFA_DRA',
        )

        usuario_avaliador = User.objects.create_user(
            username='avaliador.integridade@ufac.br',
            password='Senha123!',
            first_name='João',
            last_name='Avaliador',
            is_active=True,
        )

        self.avaliador = pUsuario.objects.create(
            usuario=usuario_avaliador,
            perfil='DOCENTE',
            titulacao='PROF_DR',
        )

        self.espaco = EspacoFisico.objects.create(
            nome='Laboratório de Integridade',
        )

        discente = Discente.objects.create(
            nome='Discente da Integridade',
            matricula='20260000123',
        )

        self.projeto = ProjetoTCC.objects.create(
            titulo='Projeto íntegro para auditoria',
            resumo='Resumo utilizado na auditoria.',
            semestre_letivo='2026.2',
            discente=discente,
            status='APROVADO',
        )

        inicio = timezone.now() + timedelta(days=20)
        fim = inicio + timedelta(hours=2)

        self.solicitacao = SolicitacaoAgendamento.objects.create(
            usuario_solicitante=self.orientador,
            projeto_tcc=self.projeto,
            espaco=self.espaco,
            opcao_data_inicio=inicio,
            opcao_data_fim=fim,
            status='APROVADA',
        )

        self.composicao = ComposicaoBanca.objects.create(
            projeto_tcc=self.projeto,
            solicitacao=self.solicitacao,
            orientador=self.orientador,
            avaliador_interno=self.avaliador,
            presidente=self.orientador,
            nome_avaliador_externo='Carlos Externo',
            titulacao_avaliador_externo='PROF_DR',
            instituicao_avaliador_externo='IFAC',
        )

        self.banca = BancaTCC.objects.create(
            solicitacao=self.solicitacao,
            projeto_tcc=self.projeto,
            espaco=self.espaco,
            data_horario_inicio=inicio,
            data_horario_fim=fim,
            status='AGENDADA',
        )


class RepositorioEDocumentoTests(BaseIntegridadeTestCase):

    def test_tela_explica_diferenca_entre_atas_e_repositorio(self):

        self.client.force_login(
            self.usuario_coordenacao
        )

        response = self.client.get(
            reverse('documentos')
        )

        self.assertContains(
            response,
            'Atas geradas automaticamente'
        )
        self.assertContains(
            response,
            'Repositório institucional de referência'
        )
        self.assertContains(
            response,
            'não alteram automaticamente'
        )

    def _dados_finalizados(self):

        self.banca.status = 'FINALIZADA'
        self.banca.nota = '9.50'

        self.banca.save(
            update_fields=[
                'status',
                'nota',
            ]
        )

        return montar_dados_ata(
            self.solicitacao,
            self.composicao,
            self.banca,
        )

    def test_documento_final_exibe_aprovacao_e_instituicao_externa(self):

        dados = self._dados_finalizados()

        documento = Document(
            gerar_docx_ata(dados)
        )

        texto = '\n'.join(
            paragrafo.text
            for paragrafo in documento.paragraphs
        )

        self.assertIn(
            'deliberou pela APROVAÇÃO',
            texto
        )
        self.assertIn(
            'Instituição: IFAC',
            texto
        )
        self.assertNotIn(
            'Instituição: UFAC',
            texto
        )

    def test_docx_possui_metadados_institucionais(self):

        documento = Document(
            gerar_docx_ata(
                self._dados_finalizados()
            )
        )

        self.assertEqual(
            documento.core_properties.author,
            'Universidade Federal do Acre'
        )
        self.assertIn(
            'Ata de Apresentação',
            documento.core_properties.title
        )


class AuditoriaIntegridadeTests(BaseIntegridadeTestCase):

    def test_base_coerente_nao_apresenta_inconsistencias(self):

        saida = StringIO()

        call_command(
            'auditar_integridade_sgtcc',
            stdout=saida,
        )

        self.assertIn(
            'Nenhuma inconsistência foi encontrada.',
            saida.getvalue()
        )

    def test_auditoria_localiza_orientador_legado_divergente(self):

        self.composicao.orientador = self.avaliador
        self.composicao.presidente = self.avaliador
        self.composicao.save(
            update_fields=[
                'orientador',
                'presidente',
            ]
        )

        saida = StringIO()

        call_command(
            'auditar_integridade_sgtcc',
            stdout=saida,
        )

        texto = saida.getvalue()

        self.assertIn(
            'ORIENTADOR_DIFERENTE_DO_SOLICITANTE',
            texto
        )
        self.assertIn(
            str(self.composicao.id),
            texto
        )
        self.assertIn(
            'Nenhum dado foi alterado.',
            texto
        )

    def test_auditoria_pode_falhar_para_uso_em_qa(self):

        self.banca.delete()

        with self.assertRaises(CommandError):
            call_command(
                'auditar_integridade_sgtcc',
                falhar_em_inconsistencia=True,
                stdout=StringIO(),
            )


class ConexaoSQLiteFalsa:

    vendor = 'sqlite'

    def __init__(self, caminho):

        self.settings_dict = {
            'NAME': str(caminho),
        }
        self.connection = sqlite3.connect(caminho)

    def ensure_connection(self):
        return None


class BackupSGTCCTests(TestCase):

    def setUp(self):
        self.temporario = TemporaryDirectory()
        self.raiz = Path(self.temporario.name)
        self.caminho_banco = self.raiz / 'origem.sqlite3'
        self.media = self.raiz / 'media'
        self.media.mkdir()

        self.conexao = ConexaoSQLiteFalsa(
            self.caminho_banco
        )

        self.conexao.connection.execute(
            'CREATE TABLE exemplo (id INTEGER PRIMARY KEY, nome TEXT)'
        )
        self.conexao.connection.execute(
            'INSERT INTO exemplo (nome) VALUES (?)',
            ('registro preservado',)
        )
        self.conexao.connection.commit()

        (self.media / 'arquivo.pdf').write_bytes(
            b'%PDF-1.4\nconteudo de teste'
        )

    def tearDown(self):
        self.conexao.connection.close()
        self.temporario.cleanup()

    def test_backup_contem_banco_midia_e_manifesto(self):

        destino = self.raiz / 'backup.zip'

        with (
            self.settings(MEDIA_ROOT=self.media),
            patch(
                'bancasapp.management.commands.'
                'criar_backup_sgtcc.connections',
                {'default': self.conexao},
            ),
        ):
            call_command(
                'criar_backup_sgtcc',
                saida=str(destino),
                stdout=StringIO(),
            )

        self.assertTrue(destino.is_file())

        with ZipFile(destino) as pacote:

            self.assertIn(
                'database/db.sqlite3',
                pacote.namelist()
            )
            self.assertIn(
                'media/arquivo.pdf',
                pacote.namelist()
            )
            self.assertIn(
                'manifesto.json',
                pacote.namelist()
            )

            manifesto = json.loads(
                pacote.read('manifesto.json')
            )

            self.assertEqual(
                manifesto['formato'],
                'backup-sgtcc-v1'
            )
            self.assertEqual(
                manifesto['midia']['quantidade_arquivos'],
                1
            )

    def test_backup_nao_sobrescreve_arquivo_existente(self):

        destino = self.raiz / 'existente.zip'
        destino.write_bytes(b'conteudo anterior')

        with (
            self.settings(MEDIA_ROOT=self.media),
            patch(
                'bancasapp.management.commands.'
                'criar_backup_sgtcc.connections',
                {'default': self.conexao},
            ),
        ):
            with self.assertRaises(CommandError):
                call_command(
                    'criar_backup_sgtcc',
                    saida=str(destino),
                    stdout=StringIO(),
                )

        self.assertEqual(
            destino.read_bytes(),
            b'conteudo anterior'
        )
