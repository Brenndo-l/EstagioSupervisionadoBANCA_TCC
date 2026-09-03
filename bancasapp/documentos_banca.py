import os
from datetime import timedelta
from decimal import Decimal
from html import escape
from io import BytesIO

import reportlab
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from reportlab.lib.enums import (
    TA_CENTER,
    TA_JUSTIFY,
    TA_LEFT,
    TA_RIGHT,
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


MESES = (
    '',
    'janeiro',
    'fevereiro',
    'março',
    'abril',
    'maio',
    'junho',
    'julho',
    'agosto',
    'setembro',
    'outubro',
    'novembro',
    'dezembro',
)


UNIDADES = {
    0: 'zero',
    1: 'um',
    2: 'dois',
    3: 'três',
    4: 'quatro',
    5: 'cinco',
    6: 'seis',
    7: 'sete',
    8: 'oito',
    9: 'nove',
    10: 'dez',
    11: 'onze',
    12: 'doze',
    13: 'treze',
    14: 'quatorze',
    15: 'quinze',
    16: 'dezesseis',
    17: 'dezessete',
    18: 'dezoito',
    19: 'dezenove',
}


DEZENAS = {
    20: 'vinte',
    30: 'trinta',
    40: 'quarenta',
    50: 'cinquenta',
    60: 'sessenta',
    70: 'setenta',
    80: 'oitenta',
    90: 'noventa',
}


def numero_ate_99_por_extenso(valor):

    if valor < 20:
        return UNIDADES[valor]

    dezena = (valor // 10) * 10
    unidade = valor % 10

    if unidade == 0:
        return DEZENAS[dezena]

    return (
        f'{DEZENAS[dezena]} e '
        f'{UNIDADES[unidade]}'
    )


def formatar_nota(nota):

    if nota is None:
        return '________________'

    nota_decimal = Decimal(nota).quantize(
        Decimal('0.01')
    )

    parte_inteira = int(nota_decimal)
    parte_decimal = int(
        (nota_decimal - parte_inteira) * 100
    )

    valor_formatado = (
        f'{nota_decimal:.2f}'
        .replace('.', ',')
    )

    valor_extenso = numero_ate_99_por_extenso(
        parte_inteira
    )

    if parte_decimal:
        valor_extenso = (
            f'{valor_extenso} vírgula '
            f'{numero_ate_99_por_extenso(parte_decimal)}'
        )

    return (
        f'{valor_formatado} '
        f'({valor_extenso})'
    )


def nome_docente(perfil):

    if perfil is None:
        return ''

    return perfil.nome_para_documento


def _funcao_com_presidencia(
    composicao,
    perfil,
    funcao
):

    presidente_id = (
        composicao.presidente_id
        or composicao.orientador_id
    )

    if (
        perfil
        and presidente_id == perfil.id
    ):
        return f'{funcao} e presidente'

    return funcao


def montar_integrantes(composicao):

    integrantes = []

    integrantes_internos = (
        (
            composicao.orientador,
            'Orientador(a)',
        ),
        (
            composicao.coorientador,
            'Coorientador(a)',
        ),
        (
            composicao.avaliador_interno,
            'Avaliador(a) interno(a)',
        ),
        (
            composicao.segundo_avaliador_interno,
            'Segundo(a) avaliador(a) interno(a)',
        ),
    )

    for perfil, funcao in integrantes_internos:

        if perfil is None:
            continue

        integrantes.append(
            {
                'nome': nome_docente(perfil),
                'funcao': _funcao_com_presidencia(
                    composicao,
                    perfil,
                    funcao
                ),
                # Nos modelos oficiais, a instituição é necessária
                # para identificar o membro externo. Para docentes da
                # UFAC, a vinculação já está indicada pela função.
                'instituicao': '',
            }
        )

    if composicao.nome_avaliador_externo:

        nome_externo = (
            composicao.nome_avaliador_externo
        )

        if composicao.titulacao_avaliador_externo:
            nome_externo = (
                f'{composicao.get_titulacao_avaliador_externo_display()} '
                f'{nome_externo}'
            )

        integrantes.append(
            {
                'nome': nome_externo,
                'funcao': 'Avaliador(a) externo(a)',
                'instituicao': (
                    composicao.instituicao_avaliador_externo
                    or 'Não informada'
                ),
            }
        )

    return integrantes


def montar_dados_ata(
    solicitacao,
    composicao,
    banca
):

    data_defesa = timezone.localtime(
        solicitacao.opcao_data_inicio
    )

    finalizada = bool(
        banca
        and banca.status == 'FINALIZADA'
        and banca.nota is not None
    )

    nota = (
        banca.nota
        if finalizada
        else None
    )

    data_limite = (
        banca.data_limite_versao_final
        if banca is not None
        else (
            data_defesa.date()
            + timedelta(days=30)
        )
    )

    return {
        'discente': solicitacao.projeto_tcc.discente.nome,
        'titulo_tcc': solicitacao.projeto_tcc.titulo,
        'espaco': solicitacao.espaco.nome,
        'presidente': nome_docente(
            composicao.presidente
            or composicao.orientador
        ),
        'integrantes': montar_integrantes(
            composicao
        ),
        'data_defesa': data_defesa,
        'data_defesa_extenso': (
            f'{data_defesa.day} de '
            f'{MESES[data_defesa.month]} de '
            f'{data_defesa.year}'
        ),
        'hora_defesa': (
            data_defesa.strftime('%Hh%M')
            if data_defesa.minute
            else data_defesa.strftime('%Hh')
        ),
        'finalizada': finalizada,
        'resultado': (
            'APROVAÇÃO'
            if finalizada
            else ''
        ),
        'nota': nota,
        'nota_formatada': formatar_nota(nota),
        'data_limite_versao_final': data_limite,
        'data_limite_extenso': (
            f'{data_limite.day} de '
            f'{MESES[data_limite.month]} de '
            f'{data_limite.year}'
        ),
    }


def _registrar_fontes_pdf():

    if 'AtaSans' in pdfmetrics.getRegisteredFontNames():
        return

    diretorio_fontes = os.path.join(
        os.path.dirname(reportlab.__file__),
        'fonts'
    )

    pdfmetrics.registerFont(
        TTFont(
            'AtaSans',
            os.path.join(diretorio_fontes, 'Vera.ttf')
        )
    )

    pdfmetrics.registerFont(
        TTFont(
            'AtaSans-Bold',
            os.path.join(diretorio_fontes, 'VeraBd.ttf')
        )
    )

    pdfmetrics.registerFontFamily(
        'AtaSans',
        normal='AtaSans',
        bold='AtaSans-Bold',
    )


def _linha_integrante_pdf(integrante):

    texto = (
        f'<b>{escape(integrante["nome"])}</b> - '
        f'{escape(integrante["funcao"])}'
    )

    if integrante['instituicao']:
        texto += (
            ' - Instituição: '
            f'{escape(integrante["instituicao"])}'
        )

    return texto


def gerar_pdf_ata(dados):

    _registrar_fontes_pdf()

    saida = BytesIO()

    documento = SimpleDocTemplate(
        saida,
        pagesize=A4,
        rightMargin=2.0 * cm,
        leftMargin=2.0 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.7 * cm,
        title=(
            'Ata de Apresentação do Trabalho '
            'de Conclusão de Curso'
        ),
        author='Universidade Federal do Acre',
    )

    estilo_cabecalho_principal = ParagraphStyle(
        'CabecalhoPrincipalAta',
        fontName='AtaSans-Bold',
        fontSize=11,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=1,
    )

    estilo_cabecalho = ParagraphStyle(
        'CabecalhoAta',
        fontName='AtaSans',
        fontSize=10.5,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=1,
    )

    estilo_titulo = ParagraphStyle(
        'TituloAta',
        fontName='AtaSans-Bold',
        fontSize=13,
        leading=16,
        alignment=TA_CENTER,
        spaceBefore=20,
        spaceAfter=18,
    )

    estilo_resumo = ParagraphStyle(
        'ResumoAta',
        fontName='AtaSans',
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        leftIndent=7.5 * cm,
        spaceAfter=22,
    )

    estilo_corpo = ParagraphStyle(
        'CorpoAta',
        fontName='AtaSans',
        fontSize=10.5,
        leading=15,
        alignment=TA_JUSTIFY,
        firstLineIndent=1.25 * cm,
        spaceAfter=10,
    )

    estilo_secao = ParagraphStyle(
        'SecaoAta',
        fontName='AtaSans-Bold',
        fontSize=10.5,
        leading=14,
        alignment=TA_CENTER,
        spaceBefore=8,
        spaceAfter=7,
    )

    estilo_integrante = ParagraphStyle(
        'IntegranteAta',
        fontName='AtaSans',
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=4,
    )

    estilo_campos = ParagraphStyle(
        'CamposAta',
        fontName='AtaSans-Bold',
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        leftIndent=1.25 * cm,
        spaceAfter=10,
    )

    estilo_data = ParagraphStyle(
        'DataAta',
        fontName='AtaSans',
        fontSize=10.5,
        leading=15,
        alignment=TA_RIGHT,
        spaceBefore=8,
    )

    elementos = [
        Paragraph(
            'UNIVERSIDADE FEDERAL DO ACRE',
            estilo_cabecalho_principal
        ),
        Paragraph(
            'Centro de Ciências Exatas e Tecnológicas',
            estilo_cabecalho
        ),
        Paragraph(
            (
                'Coordenação do Curso de Bacharelado '
                'em Sistemas de Informação'
            ),
            estilo_cabecalho
        ),
        Paragraph(
            (
                'ATA DE APRESENTAÇÃO DO TRABALHO '
                'DE CONCLUSÃO DE CURSO'
            ),
            estilo_titulo
        ),
        Paragraph(
            (
                'Ata de apresentação do Trabalho de Conclusão '
                'do Curso de Bacharelado em Sistemas de '
                'Informação, realizada no dia '
                f'<b>{escape(dados["data_defesa_extenso"])}</b>.'
            ),
            estilo_resumo
        ),
    ]

    texto_corpo = (
        'No dia '
        f'<b>{escape(dados["data_defesa_extenso"])}</b>, às '
        f'<b>{escape(dados["hora_defesa"])}</b>, no espaço '
        f'{escape(dados["espaco"])}, desta Universidade e na '
        'presença da Banca Examinadora presidida por '
        f'{escape(dados["presidente"])} e composta pelos '
        'membros relacionados abaixo, o(a) discente '
        f'<b>{escape(dados["discente"])}</b> realizou a Defesa '
        'Pública do Trabalho de Conclusão de Curso, intitulado '
        f'<b>“{escape(dados["titulo_tcc"])}”</b>, como requisito '
        'curricular indispensável à integralização do Curso de '
        'Bacharelado em Sistemas de Informação.'
    )

    elementos.append(
        Paragraph(texto_corpo, estilo_corpo)
    )

    elementos.append(
        Paragraph(
            'Banca Examinadora',
            estilo_secao
        )
    )

    for integrante in dados['integrantes']:
        elementos.append(
            Paragraph(
                _linha_integrante_pdf(integrante),
                estilo_integrante
            )
        )

    elementos.append(Spacer(1, 9))

    if dados['finalizada']:
        texto_resultado = (
            'A Banca Examinadora, após reunião em sessão '
            'reservada, deliberou pela '
            f'<b>{escape(dados["resultado"])}</b> do referido '
            'Trabalho de Conclusão de Curso e atribuiu a nota '
            f'<b>{escape(dados["nota_formatada"])}</b>, '
            'divulgando o resultado formalmente ao(à) discente '
            'e aos demais presentes.'
        )
    else:
        texto_resultado = (
            'A Banca Examinadora, após reunião em sessão '
            'reservada, registrará o resultado da avaliação '
            'e a nota final nos campos abaixo.'
        )

    elementos.append(
        Paragraph(texto_resultado, estilo_corpo)
    )

    if not dados['finalizada']:
        elementos.append(
            Paragraph(
                (
                    'Resultado: ______________________________'
                    '&nbsp;&nbsp;&nbsp;&nbsp;'
                    'Nota: __________________'
                ),
                estilo_campos
            )
        )

    if dados['finalizada']:
        texto_prazo = (
            'Ao final, o(a) discente foi informado(a) da '
            'obrigatoriedade da apresentação da versão final '
            'do Trabalho de Conclusão de Curso no prazo máximo '
            'de <b>30 (trinta) dias corridos</b>, até '
            f'<b>{escape(dados["data_limite_extenso"])}</b>, '
            'contendo todos os ajustes sugeridos pela Banca '
            'Examinadora, sob consentimento do(a) orientador(a).'
        )
    else:
        texto_prazo = (
            'Após a defesa, o(a) discente deverá ser '
            'informado(a) da obrigatoriedade da apresentação '
            'da versão final do Trabalho de Conclusão de Curso '
            'no prazo máximo de <b>30 (trinta) dias corridos</b>, '
            'até '
            f'<b>{escape(dados["data_limite_extenso"])}</b>, '
            'contendo todos os ajustes sugeridos pela Banca '
            'Examinadora, sob consentimento do(a) orientador(a).'
        )

    elementos.extend(
        [
            Paragraph(texto_prazo, estilo_corpo),
            Paragraph(
                'Por ser verdade, firmamos o presente.',
                estilo_corpo
            ),
            Paragraph(
                (
                    'Rio Branco-AC, '
                    f'{escape(dados["data_defesa_extenso"])}.'
                ),
                estilo_data
            ),
        ]
    )

    documento.build(elementos)
    saida.seek(0)

    return saida


def _configurar_fonte_run(
    run,
    tamanho=11,
    negrito=False
):

    run.font.name = 'Times New Roman'
    run.font.size = Pt(tamanho)
    run.font.bold = negrito

    run._element.rPr.rFonts.set(
        qn('w:eastAsia'),
        'Times New Roman'
    )


def _adicionar_run(
    paragrafo,
    texto,
    tamanho=11,
    negrito=False
):

    run = paragrafo.add_run(texto)

    _configurar_fonte_run(
        run,
        tamanho=tamanho,
        negrito=negrito
    )

    return run


def _adicionar_paragrafo_centralizado(
    documento,
    texto,
    tamanho=11,
    negrito=False,
    espaco_depois=0
):

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_after = Pt(
        espaco_depois
    )

    _adicionar_run(
        paragrafo,
        texto,
        tamanho=tamanho,
        negrito=negrito
    )

    return paragrafo


def _novo_paragrafo_corpo(
    documento,
    espaco_depois=10
):

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.first_line_indent = Cm(1.25)
    paragrafo.paragraph_format.line_spacing = 1.15
    paragrafo.paragraph_format.space_after = Pt(
        espaco_depois
    )

    return paragrafo


def _adicionar_linha_integrante_docx(
    documento,
    integrante
):

    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_after = Pt(3)

    _adicionar_run(
        paragrafo,
        integrante['nome'],
        tamanho=10.5,
        negrito=True
    )

    _adicionar_run(
        paragrafo,
        f' - {integrante["funcao"]}',
        tamanho=10.5
    )

    if integrante['instituicao']:
        _adicionar_run(
            paragrafo,
            (
                ' - Instituição: '
                f'{integrante["instituicao"]}'
            ),
            tamanho=10.5
        )


def gerar_docx_ata(dados):

    documento = Document()

    documento.core_properties.title = (
        'Ata de Apresentação do Trabalho de Conclusão de Curso'
    )
    documento.core_properties.subject = (
        'Registro institucional de banca de TCC'
    )
    documento.core_properties.author = (
        'Universidade Federal do Acre'
    )
    documento.core_properties.keywords = (
        'UFAC; SGTCC; banca de TCC; ata de apresentação'
    )

    secao = documento.sections[0]
    secao.page_width = Mm(210)
    secao.page_height = Mm(297)
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.7)
    secao.left_margin = Cm(2.0)
    secao.right_margin = Cm(2.0)

    estilo_normal = documento.styles['Normal']
    estilo_normal.font.name = 'Times New Roman'
    estilo_normal.font.size = Pt(11)
    estilo_normal._element.rPr.rFonts.set(
        qn('w:eastAsia'),
        'Times New Roman'
    )

    _adicionar_paragrafo_centralizado(
        documento,
        'UNIVERSIDADE FEDERAL DO ACRE',
        tamanho=11,
        negrito=True
    )

    _adicionar_paragrafo_centralizado(
        documento,
        'Centro de Ciências Exatas e Tecnológicas',
        tamanho=10.5
    )

    _adicionar_paragrafo_centralizado(
        documento,
        (
            'Coordenação do Curso de Bacharelado '
            'em Sistemas de Informação'
        ),
        tamanho=10.5,
        espaco_depois=18
    )

    _adicionar_paragrafo_centralizado(
        documento,
        (
            'ATA DE APRESENTAÇÃO DO TRABALHO '
            'DE CONCLUSÃO DE CURSO'
        ),
        tamanho=13,
        negrito=True,
        espaco_depois=14
    )

    resumo = documento.add_paragraph()
    resumo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    resumo.paragraph_format.left_indent = Cm(7.5)
    resumo.paragraph_format.space_after = Pt(18)

    _adicionar_run(
        resumo,
        (
            'Ata de apresentação do Trabalho de Conclusão '
            'do Curso de Bacharelado em Sistemas de '
            'Informação, realizada no dia '
        ),
        tamanho=10.5
    )

    _adicionar_run(
        resumo,
        f'{dados["data_defesa_extenso"]}.',
        tamanho=10.5,
        negrito=True
    )

    corpo = _novo_paragrafo_corpo(
        documento,
        espaco_depois=10
    )

    _adicionar_run(corpo, 'No dia ')

    _adicionar_run(
        corpo,
        dados['data_defesa_extenso'],
        negrito=True
    )

    _adicionar_run(corpo, ', às ')

    _adicionar_run(
        corpo,
        dados['hora_defesa'],
        negrito=True
    )

    _adicionar_run(
        corpo,
        (
            f', no espaço {dados["espaco"]}, desta '
            'Universidade e na presença da Banca '
            'Examinadora presidida por '
            f'{dados["presidente"]} e composta pelos '
            'membros relacionados abaixo, o(a) discente '
        )
    )

    _adicionar_run(
        corpo,
        dados['discente'],
        negrito=True
    )

    _adicionar_run(
        corpo,
        (
            ' realizou a Defesa Pública do Trabalho de '
            'Conclusão de Curso, intitulado '
        )
    )

    _adicionar_run(
        corpo,
        f'“{dados["titulo_tcc"]}”',
        negrito=True
    )

    _adicionar_run(
        corpo,
        (
            ', como requisito curricular indispensável à '
            'integralização do Curso de Bacharelado em '
            'Sistemas de Informação.'
        )
    )

    _adicionar_paragrafo_centralizado(
        documento,
        'Banca Examinadora',
        tamanho=11,
        negrito=True,
        espaco_depois=5
    )

    for integrante in dados['integrantes']:
        _adicionar_linha_integrante_docx(
            documento,
            integrante
        )

    resultado = _novo_paragrafo_corpo(
        documento,
        espaco_depois=10
    )
    resultado.paragraph_format.space_before = Pt(9)

    if dados['finalizada']:
        _adicionar_run(
            resultado,
            (
                'A Banca Examinadora, após reunião em sessão '
                'reservada, deliberou pela '
            )
        )

        _adicionar_run(
            resultado,
            dados['resultado'],
            negrito=True
        )

        _adicionar_run(
            resultado,
            (
                ' do referido Trabalho de Conclusão de Curso '
                'e atribuiu a nota '
            )
        )

        _adicionar_run(
            resultado,
            dados['nota_formatada'],
            negrito=True
        )

        _adicionar_run(
            resultado,
            (
                ', divulgando o resultado formalmente ao(à) '
                'discente e aos demais presentes.'
            )
        )

    else:
        _adicionar_run(
            resultado,
            (
                'A Banca Examinadora, após reunião em sessão '
                'reservada, registrará o resultado da avaliação '
                'e a nota final nos campos abaixo.'
            )
        )

        campos = documento.add_paragraph()
        campos.alignment = WD_ALIGN_PARAGRAPH.LEFT
        campos.paragraph_format.left_indent = Cm(1.25)
        campos.paragraph_format.space_after = Pt(10)

        _adicionar_run(
            campos,
            (
                'Resultado: ______________________________    '
                'Nota: __________________'
            ),
            negrito=True
        )

    prazo = _novo_paragrafo_corpo(
        documento,
        espaco_depois=10
    )

    if dados['finalizada']:
        _adicionar_run(
            prazo,
            (
                'Ao final, o(a) discente foi informado(a) da '
                'obrigatoriedade da apresentação da versão '
                'final do Trabalho de Conclusão de Curso no '
                'prazo máximo de '
            )
        )
    else:
        _adicionar_run(
            prazo,
            (
                'Após a defesa, o(a) discente deverá ser '
                'informado(a) da obrigatoriedade da apresentação '
                'da versão final do Trabalho de Conclusão de '
                'Curso no prazo máximo de '
            )
        )

    _adicionar_run(
        prazo,
        '30 (trinta) dias corridos',
        negrito=True
    )

    _adicionar_run(prazo, ', até ')

    _adicionar_run(
        prazo,
        dados['data_limite_extenso'],
        negrito=True
    )

    _adicionar_run(
        prazo,
        (
            ', contendo todos os ajustes sugeridos pela '
            'Banca Examinadora, sob consentimento do(a) '
            'orientador(a).'
        )
    )

    encerramento = _novo_paragrafo_corpo(
        documento,
        espaco_depois=16
    )

    _adicionar_run(
        encerramento,
        'Por ser verdade, firmamos o presente.'
    )

    data = documento.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _adicionar_run(
        data,
        f'Rio Branco-AC, {dados["data_defesa_extenso"]}.'
    )

    saida = BytesIO()
    documento.save(saida)
    saida.seek(0)

    return saida
